import os
import PyPDF2
from docx import Document
from src.nlp_processing import extract_skills_from_description
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from crewai import Agent  # NEW: Import for crewAI
from transformers import pipeline  # NEW: Import for LLM

# NEW: Load BART model
# generator = pipeline("text2text-generation", model="facebook/bart-large")
# Global model to avoid reloading
if "generator" not in globals():
    generator = pipeline("text2text-generation", model="sshleifer/distilbart-cnn-12-6")  # Smaller model for speed


# def generate_cover_letter(job_title, company, job_desc, cv_file_path):
#     # REMOVED: Old implementation using extract_skills_from_description
#     # Extract skills from job description
#     skills = extract_skills_from_description(job_desc)
    
#     # Extract relevant experience from CV
#     experience = extract_experience_from_cv(cv_file_path)

#     # Extract name and contact info from the CV
#     name, contact_info = extract_name_and_contact_from_cv(cv_file_path)
    
#     # NEW: Use LLM instead of static template
#     prompt = (
#         f"Write a professional cover letter for the {job_title} position at {company}. "
#         f"Job description: {job_desc[:200]}. Skills: {', '.join(skills)}. "
#         f"My experience: {experience[:200]}. Name: {name}. Contact: {contact_info}. "
#         "Keep it concise, under 300 words."
#     )
#     result = generator(prompt, max_length=300, num_return_sequences=1)
#     return result[0]["generated_text"]

    # # Create the cover letter template
    # cover_letter = f"""
    # Dear Hiring Manager,

    # I am excited to apply for the {job_title} position at {company}. With a strong background in {', '.join(skills)}, 
    # I am eager to contribute my expertise to your team.

    # Currently, I am a {experience}. I have successfully contributed to optimizing experimental workflows, improving predictive accuracy, and co-authoring research papers. Furthermore, I have disseminated AI/ML research publications to a community of over 200 data scientists, helping foster collaboration and drive advancements in the field.

    # I am excited about the opportunity to apply my technical skills, analytical expertise, and academic background to the {job_title} role at {company}, where I can leverage my experience to help drive positive change in the advertising industry. The opportunity to work on flexible hours and contribute to a team that ensures responsible advertising resonates with my professional values and long-term career goals.

    # I look forward to the opportunity to further discuss how my background and skills can contribute to the continued success of {company}. Thank you for your consideration.

    # Sincerely,  
    # {name}  
    # {contact_info}
    # """
    
    # return cover_letter

# Helper functions for extracting information from CV
def extract_experience_from_cv(cv_file_path):
    experience = ""

    if cv_file_path.lower().endswith('.pdf'):
        with open(cv_file_path, 'rb') as cv_file:
            reader = PyPDF2.PdfReader(cv_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        experience = extract_experience_from_text(text)

    elif cv_file_path.lower().endswith('.docx'):
        doc = Document(cv_file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        experience = extract_experience_from_text(text)

    return experience

def extract_experience_from_text(text):
    experience_section = ""
    experience_keywords = ['experience', 'work', 'role', 'responsibilities']

    for line in text.split('\n'):
        for keyword in experience_keywords:
            if keyword in line.lower():
                experience_section += line.strip() + "\n"
                break

    return experience_section

def extract_name_and_contact_from_cv(cv_file_path):
    name = ""
    contact_info = ""

    if cv_file_path.lower().endswith('.pdf'):
        with open(cv_file_path, 'rb') as cv_file:
            reader = PyPDF2.PdfReader(cv_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        name, contact_info = extract_name_and_contact_from_text(text)

    elif cv_file_path.lower().endswith('.docx'):
        doc = Document(cv_file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        name, contact_info = extract_name_and_contact_from_text(text)

    return name, contact_info

def extract_name_and_contact_from_text(text):
    lines = text.split("\n")
    name = lines[0] if len(lines) > 0 else "Your Full Name"
    contact_info = lines[1] if len(lines) > 1 else "Your Contact Information"
    
    return name, contact_info

# Save the CV and Cover Letter to Files
# def save_to_files(cv_file, cover_letter, name):
#     output_dir = "generated_documents"
#     if not os.path.exists(output_dir):
#         os.makedirs(output_dir)

#     cover_letter_filename = f"Cover_letter_{name}.txt"
#     cv_filename = f"CV_{name}.txt"
    
#     with open(os.path.join(output_dir, cover_letter_filename), "w") as f:
#         f.write(cover_letter)
    
#     with open(os.path.join(output_dir, cv_filename), "w") as f:
#         f.write("CV content goes here...")  # This should ideally be CV content extracted from the file

#     return os.path.join(output_dir, cover_letter_filename), os.path.join(output_dir, cv_filename)
def save_to_files(cv_file, cover_letter, name):
    """Save cover letter to disk."""
    output_dir = "generated_documents"
    os.makedirs(output_dir, exist_ok=True)
    cover_letter_path = os.path.join(output_dir, f"Cover_letter_{name}.txt")
    with open(cover_letter_path, "w") as f:
        f.write(cover_letter)
 
    return cover_letter_path, None  # Simplified to return only cover letter path
# Function to send email with CV and Cover Letter as attachments
def send_email(subject, body, recipient, cv_path, cover_letter_path):
    msg = MIMEMultipart()
    msg['From'] = 'youremail@example.com'
    msg['To'] = recipient
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        with open(cv_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cv_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching CV: {e}")

    try:
        with open(cover_letter_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cover_letter_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching cover letter: {e}")

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login('youremail@example.com', '1891')
        server.sendmail(msg['From'], recipient, msg.as_string())
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"Error sending email: {e}")
        
from transformers import BartForConditionalGeneration, BartTokenizer
# Load the model and tokenizer
model_name = "facebook/bart-large"
tokenizer = BartTokenizer.from_pretrained(model_name)
model = BartForConditionalGeneration.from_pretrained(model_name)

def generate_cover_letter(job_title, company, job_description, user_cv):
    input_text = f"Write a cover letter for the position of {job_title} at {company}. Job description: {job_description}. My CV: {user_cv}"
    inputs = tokenizer(input_text, return_tensors="pt", max_length=512, truncation=True)
    outputs = model.generate(**inputs, max_length=300, num_beams=4, early_stopping=True)
    cover_letter = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return cover_letter       
# NEW: Define Writer Agent
writer_agent = Agent(
    role="Cover Letter Writer",
    goal="Generate tailored cover letters for job applications",
    backstory="Skilled in crafting professional letters",
    verbose=False,
    tools=[generate_cover_letter]
)
