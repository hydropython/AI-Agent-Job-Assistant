import os
import PyPDF2
from docx import Document
from nlp_processing import extract_skills_from_description
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

def generate_cover_letter(job_title, company, job_desc, cv_file):
    # Extract skills from job description
    skills = extract_skills_from_description(job_desc)
    
    # Extract relevant experience from CV
    experience = extract_experience_from_cv(cv_file)

    # Extract name and contact info from the CV
    name, contact_info = extract_name_and_contact_from_cv(cv_file)

    # Create the cover letter template
    cover_letter = f"""
    Dear Hiring Manager,

    I am excited to apply for the {job_title} position at {company}. With a strong background in {', '.join(skills)}, 
    I am eager to contribute my expertise to your team.

    Currently, I am a {experience}. I have successfully contributed to optimizing experimental workflows, improving predictive accuracy, and co-authoring research papers. Furthermore, I have disseminated AI/ML research publications to a community of over 200 data scientists, helping foster collaboration and drive advancements in the field.

    I am excited about the opportunity to apply my technical skills, analytical expertise, and academic background to the {job_title} role at {company}, where I can leverage my experience to help drive positive change in the advertising industry. The opportunity to work on flexible hours and contribute to a team that ensures responsible advertising resonates with my professional values and long-term career goals.

    I look forward to the opportunity to further discuss how my background and skills can contribute to the continued success of {company}. Thank you for your consideration.

    Sincerely,  
    {name}  
    {contact_info}
    """
    
    return cover_letter

# Helper functions for extracting information from CV
def extract_experience_from_cv(cv_file):
    experience = ""

    if cv_file.name.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(cv_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        experience = extract_experience_from_text(text)

    elif cv_file.name.lower().endswith('.docx'):
        doc = Document(cv_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        experience = extract_experience_from_text(text)

    return experience

def extract_experience_from_text(text):
    experience_section = ""
    experience_keywords = ['experience', 'work', 'role', 'responsibilities']

    for line in text.split('\n'):
        for keyword in experience_keywords:
            if keyword in line.lower():
                experience_section += line.strip() + "\n"
                break

    return experience_section

def extract_name_and_contact_from_cv(cv_file):
    name = ""
    contact_info = ""

    if cv_file.name.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(cv_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        name, contact_info = extract_name_and_contact_from_text(text)

    elif cv_file.name.lower().endswith('.docx'):
        doc = Document(cv_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        name, contact_info = extract_name_and_contact_from_text(text)

    return name, contact_info

def extract_name_and_contact_from_text(text):
    lines = text.split("\n")
    name = lines[0] if len(lines) > 0 else "Your Full Name"
    contact_info = lines[1] if len(lines) > 1 else "Your Contact Information"
    
    return name, contact_info

# Save the CV and Cover Letter to Files
def save_to_files(cv_file, cover_letter, name):
    output_dir = "generated_documents"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    cover_letter_filename = f"Cover_letter_{name}.txt"
    cv_filename = f"CV_{name}.txt"
    
    with open(os.path.join(output_dir, cover_letter_filename), "w") as f:
        f.write(cover_letter)
    
    with open(os.path.join(output_dir, cv_filename), "w") as f:
        f.write("CV content goes here...")  # This should ideally be CV content extracted from the file

    return os.path.join(output_dir, cover_letter_filename), os.path.join(output_dir, cv_filename)

# Function to send email with CV and Cover Letter as attachments
def send_email(subject, body, recipient, cv_path, cover_letter_path):
    msg = MIMEMultipart()
    msg['From'] = 'sender@example.com'
    msg['To'] = recipient
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        with open(cv_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cv_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching CV: {e}")

    try:
        with open(cover_letter_path, 'rb') as file:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(cover_letter_path)}')
            msg.attach(part)
    except Exception as e:
        print(f"Error attaching cover letter: {e}")

    try:
        server = smtplib.SMTP('smtp.example.com', 587)
        server.starttls()
        server.login('your_email@example.com', 'your_password')
        server.sendmail(msg['From'], recipient, msg.as_string())
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"Error sending email: {e}")

def main():
    job_title = "Data Scientist"
    company = "Tech Company"
    job_desc = "We are looking for a Data Scientist with strong skills in machine learning and data analysis."
    cv_path = 'path_to_cv.pdf'  # Replace with actual CV path
    
    with open(cv_path, "rb") as cv_file:
        cover_letter = generate_cover_letter(job_title, company, job_desc, cv_file)
        name, contact_info = extract_name_and_contact_from_cv(cv_file)
        cover_letter_filename, cv_filename = save_to_files(cv_file, cover_letter, name)

        # Send the email with attachments
        send_email(
            subject="Application for Data Scientist",
            body="Please find my application attached.",
            recipient="recipient@example.com",  # Change this to the actual recipient
            cv_path=cover_letter_filename,
            cover_letter_path=cv_filename
        )

if __name__ == "__main__":
    main()

